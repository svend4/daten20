#!/usr/bin/env python3
"""
DOCX Export Examples

Demonstrates various ways to use the DOCX exporter:
1. Simple text export
2. Structured data export
3. Professional report
4. Custom branding
5. Tables and lists
6. Images and formatting
7. Multi-section document
8. Invoice/Quote document
"""

import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.core.docx_exporter import DOCXExporter, BrandingConfig, export_to_docx
from datetime import datetime


def example_1_simple_text():
    """Example 1: Simple text export."""
    print("\n" + "="*70)
    print("Example 1: Simple Text Export")
    print("="*70)

    content = """
# Introduction

This is a simple document created with the DOCX exporter.

## Features

The exporter supports:
- Automatic heading detection
- Bullet lists
- Numbered lists
- Custom formatting

## Conclusion

This demonstrates the basic functionality of the DOCX exporter.
"""

    output_path = "data/exports/example1_simple.docx"
    success = export_to_docx(content, output_path, title="Simple Document")

    if success:
        print(f"✓ Document created: {output_path}")
    else:
        print(f"✗ Failed to create document")


def example_2_structured_data():
    """Example 2: Export structured data."""
    print("\n" + "="*70)
    print("Example 2: Structured Data Export")
    print("="*70)

    data = {
        "company_info": {
            "name": "Acme Corporation",
            "address": "123 Main Street, Berlin, Germany",
            "phone": "+49 30 12345678",
            "email": "info@acme.com"
        },
        "employees": [
            {"name": "John Doe", "position": "CEO", "salary": "120000"},
            {"name": "Jane Smith", "position": "CTO", "salary": "100000"},
            {"name": "Bob Johnson", "position": "Developer", "salary": "80000"}
        ],
        "departments": {
            "Engineering": "Responsible for product development",
            "Sales": "Responsible for customer acquisition",
            "HR": "Responsible for employee management"
        },
        "summary": "This is a structured document showing company information, employee list, and departments."
    }

    output_path = "data/exports/example2_structured.docx"
    success = export_to_docx(data, output_path, title="Company Information")

    if success:
        print(f"✓ Document created: {output_path}")
    else:
        print(f"✗ Failed to create document")


def example_3_professional_report():
    """Example 3: Professional report with sections."""
    print("\n" + "="*70)
    print("Example 3: Professional Report")
    print("="*70)

    report_data = {
        "title": "Quarterly Business Report",
        "subtitle": "Q4 2025 Performance Analysis",
        "date": datetime.now().strftime("%Y-%m-%d"),
        "summary": "This quarter showed strong performance across all key metrics. Revenue increased by 25%, customer satisfaction improved to 92%, and we successfully launched 3 new products.",
        "sections": [
            {
                "title": "Financial Performance",
                "content": "Revenue for Q4 2025 reached €2.5M, representing a 25% increase over Q3.",
                "subsections": [
                    {
                        "title": "Revenue Breakdown",
                        "content": "Revenue came from multiple sources:",
                        "table": {
                            "headers": ["Source", "Amount", "Percentage"],
                            "rows": [
                                ["Product Sales", "€1,500,000", "60%"],
                                ["Services", "€750,000", "30%"],
                                ["Licensing", "€250,000", "10%"]
                            ]
                        }
                    },
                    {
                        "title": "Expenses",
                        "content": "Total expenses were controlled at €1.8M:",
                        "list": [
                            "Personnel: €1,200,000 (67%)",
                            "Infrastructure: €400,000 (22%)",
                            "Marketing: €200,000 (11%)"
                        ]
                    }
                ]
            },
            {
                "title": "Customer Metrics",
                "content": "Customer satisfaction and retention both improved significantly.",
                "subsections": [
                    {
                        "title": "Key Metrics",
                        "table": {
                            "headers": ["Metric", "Q3 2025", "Q4 2025", "Change"],
                            "rows": [
                                ["Satisfaction Score", "85%", "92%", "+7%"],
                                ["Retention Rate", "88%", "91%", "+3%"],
                                ["NPS", "45", "52", "+7"]
                            ]
                        }
                    }
                ]
            },
            {
                "title": "Product Launches",
                "content": "Three new products were successfully launched:",
                "subsections": [
                    {
                        "title": "New Products",
                        "list": [
                            "Product Alpha - Document AI Assistant (October 2025)",
                            "Product Beta - Advanced Analytics Dashboard (November 2025)",
                            "Product Gamma - Mobile Application Suite (December 2025)"
                        ]
                    }
                ]
            }
        ],
        "conclusion": "Q4 2025 was an exceptional quarter. We exceeded all targets and set a strong foundation for 2026. Key priorities for Q1 2026 include expanding to new markets and scaling our infrastructure."
    }

    exporter = DOCXExporter()
    output_path = "data/exports/example3_report.docx"
    success = exporter.export_report(report_data, output_path)

    if success:
        print(f"✓ Report created: {output_path}")
    else:
        print(f"✗ Failed to create report")


def example_4_custom_branding():
    """Example 4: Document with custom branding."""
    print("\n" + "="*70)
    print("Example 4: Custom Branding")
    print("="*70)

    # Custom branding
    branding = BrandingConfig()
    branding.company_name = "TechCorp International"
    branding.company_tagline = "Innovation Through Technology"
    branding.primary_color = RGBColor(220, 38, 38)  # Red
    branding.secondary_color = RGBColor(107, 114, 128)  # Gray
    branding.accent_color = RGBColor(34, 197, 94)  # Green
    branding.footer_text = "© 2025 TechCorp International - Confidential"

    # Create document
    exporter = DOCXExporter(branding)
    exporter.create_document("Custom Brand Document")
    exporter.add_title("TechCorp International")

    exporter.add_heading("About Us", level=1)
    exporter.add_paragraph(
        "TechCorp International is a leading provider of innovative technology solutions. "
        "We specialize in AI, cloud computing, and digital transformation."
    )

    exporter.add_heading("Our Services", level=1)
    exporter.add_bullet_list([
        "Artificial Intelligence & Machine Learning",
        "Cloud Infrastructure & DevOps",
        "Digital Transformation Consulting",
        "Custom Software Development"
    ])

    exporter.add_heading("Contact Information", level=1)
    exporter.add_paragraph("For more information, please contact us:")
    exporter.add_paragraph("Email: info@techcorp-intl.com")
    exporter.add_paragraph("Phone: +1 555 123 4567")

    output_path = "data/exports/example4_branding.docx"
    exporter.doc.save(output_path)
    print(f"✓ Document with custom branding created: {output_path}")


def example_5_tables_and_lists():
    """Example 5: Complex tables and lists."""
    print("\n" + "="*70)
    print("Example 5: Tables and Lists")
    print("="*70)

    exporter = DOCXExporter()
    exporter.create_document("Tables and Lists Demo")
    exporter.add_title("Complex Data Presentation")

    # Employee table
    exporter.add_heading("Employee Directory", level=1)
    exporter.add_table(
        data=[
            ["EMP001", "Alice Johnson", "Engineering", "alice@company.com", "$95,000"],
            ["EMP002", "Bob Smith", "Sales", "bob@company.com", "$85,000"],
            ["EMP003", "Carol White", "Marketing", "carol@company.com", "$75,000"],
            ["EMP004", "David Brown", "Engineering", "david@company.com", "$90,000"]
        ],
        headers=["ID", "Name", "Department", "Email", "Salary"]
    )

    # Numbered list
    exporter.add_heading("Project Phases", level=1)
    exporter.add_numbered_list([
        "Requirements gathering and analysis",
        "System design and architecture",
        "Implementation and development",
        "Testing and quality assurance",
        "Deployment and monitoring",
        "Maintenance and support"
    ])

    # Bullet list
    exporter.add_heading("Key Technologies", level=1)
    exporter.add_bullet_list([
        "Python 3.11+",
        "Flask web framework",
        "PostgreSQL database",
        "Redis caching",
        "Docker containers",
        "Kubernetes orchestration"
    ])

    output_path = "data/exports/example5_tables_lists.docx"
    exporter.doc.save(output_path)
    print(f"✓ Document with tables and lists created: {output_path}")


def example_6_info_boxes():
    """Example 6: Info boxes and highlights."""
    print("\n" + "="*70)
    print("Example 6: Info Boxes")
    print("="*70)

    exporter = DOCXExporter()
    exporter.create_document("Info Boxes Demo")
    exporter.add_title("Document with Highlights")

    exporter.add_heading("Important Information", level=1)

    exporter.add_info_box(
        "Security Notice",
        "This document contains confidential information. Please do not share without authorization."
    )

    exporter.add_paragraph(
        "Regular paragraph text follows the info box. This demonstrates how different "
        "elements can be combined in a professional document."
    )

    exporter.add_info_box(
        "Performance Tip",
        "For optimal performance, ensure that your system has at least 8GB RAM and "
        "a modern multi-core processor."
    )

    exporter.add_heading("Best Practices", level=1)
    exporter.add_paragraph(
        "Follow these guidelines for best results with the document management system."
    )

    exporter.add_info_box(
        "Quick Tip",
        "Use keyboard shortcuts to improve your productivity: Ctrl+S to save, "
        "Ctrl+P to print, Ctrl+F to find."
    )

    output_path = "data/exports/example6_info_boxes.docx"
    exporter.doc.save(output_path)
    print(f"✓ Document with info boxes created: {output_path}")


def example_7_invoice():
    """Example 7: Professional invoice document."""
    print("\n" + "="*70)
    print("Example 7: Invoice Document")
    print("="*70)

    exporter = DOCXExporter()
    exporter.create_document("Invoice")

    # Invoice header
    exporter.add_title("INVOICE")

    # Company info
    exporter.add_paragraph("Acme Corporation")
    exporter.add_paragraph("123 Business Street")
    exporter.add_paragraph("Berlin, Germany 10115")
    exporter.add_paragraph("VAT ID: DE123456789")

    # Invoice details
    exporter.add_heading("Invoice Details", level=1)
    details_table = [
        ["Invoice Number:", "INV-2025-001"],
        ["Invoice Date:", datetime.now().strftime("%Y-%m-%d")],
        ["Due Date:", "2025-02-12"],
        ["Payment Terms:", "Net 30"]
    ]
    exporter.add_table(details_table)

    # Bill to
    exporter.add_heading("Bill To", level=1)
    exporter.add_paragraph("Customer Company GmbH")
    exporter.add_paragraph("456 Client Avenue")
    exporter.add_paragraph("Munich, Germany 80331")

    # Line items
    exporter.add_heading("Line Items", level=1)
    line_items = [
        ["1", "Professional Services - Consulting", "40", "€150.00", "€6,000.00"],
        ["2", "Software License - Annual", "1", "€2,500.00", "€2,500.00"],
        ["3", "Support & Maintenance", "12", "€200.00", "€2,400.00"],
        ["", "", "", "Subtotal:", "€10,900.00"],
        ["", "", "", "VAT (19%):", "€2,071.00"],
        ["", "", "", "Total:", "€12,971.00"]
    ]
    exporter.add_table(
        line_items,
        headers=["#", "Description", "Qty", "Unit Price", "Total"]
    )

    # Payment info
    exporter.add_heading("Payment Information", level=1)
    exporter.add_paragraph("Bank: Deutsche Bank AG")
    exporter.add_paragraph("IBAN: DE89 3704 0044 0532 0130 00")
    exporter.add_paragraph("BIC: DEUTDEDBBER")

    # Footer note
    exporter.add_info_box(
        "Note",
        "Payment is due within 30 days. Late payments may incur interest charges. "
        "Thank you for your business!"
    )

    output_path = "data/exports/example7_invoice.docx"
    exporter.doc.save(output_path)
    print(f"✓ Invoice document created: {output_path}")


def example_8_multi_section():
    """Example 8: Multi-section document with page breaks."""
    print("\n" + "="*70)
    print("Example 8: Multi-Section Document")
    print("="*70)

    exporter = DOCXExporter()
    exporter.create_document("Technical Documentation")
    exporter.add_title("Product Technical Documentation")

    # Section 1: Overview
    exporter.add_heading("1. Product Overview", level=1)
    exporter.add_paragraph(
        "This document provides comprehensive technical documentation for our "
        "document management system. It covers installation, configuration, "
        "usage, and troubleshooting."
    )

    exporter.add_page_break()

    # Section 2: Installation
    exporter.add_heading("2. Installation", level=1)
    exporter.add_heading("2.1 System Requirements", level=2)
    exporter.add_bullet_list([
        "Python 3.11 or higher",
        "8GB RAM minimum (16GB recommended)",
        "10GB free disk space",
        "Modern web browser"
    ])

    exporter.add_heading("2.2 Installation Steps", level=2)
    exporter.add_numbered_list([
        "Download the installation package",
        "Extract to desired location",
        "Run setup.py to initialize",
        "Configure environment variables",
        "Start the application"
    ])

    exporter.add_page_break()

    # Section 3: Configuration
    exporter.add_heading("3. Configuration", level=1)
    exporter.add_paragraph(
        "Configuration is managed through the config.yaml file located in the "
        "config directory. Below are the most important settings:"
    )

    config_table = [
        ["database.host", "localhost", "Database server hostname"],
        ["database.port", "5432", "Database port"],
        ["app.debug", "false", "Enable debug mode"],
        ["app.workers", "4", "Number of worker processes"]
    ]
    exporter.add_table(
        config_table,
        headers=["Setting", "Default", "Description"]
    )

    exporter.add_page_break()

    # Section 4: Troubleshooting
    exporter.add_heading("4. Troubleshooting", level=1)

    exporter.add_info_box(
        "Common Issue #1",
        "If the application fails to start, check that all dependencies are installed "
        "and the database is running."
    )

    exporter.add_info_box(
        "Common Issue #2",
        "If you see import errors, ensure that your PYTHONPATH is correctly set and "
        "all required packages are installed."
    )

    output_path = "data/exports/example8_multi_section.docx"
    exporter.doc.save(output_path)
    print(f"✓ Multi-section document created: {output_path}")


def run_all_examples():
    """Run all examples."""
    print("\n" + "="*70)
    print("DOCX EXPORT EXAMPLES")
    print("="*70)

    # Ensure output directory exists
    Path("data/exports").mkdir(parents=True, exist_ok=True)

    examples = [
        example_1_simple_text,
        example_2_structured_data,
        example_3_professional_report,
        example_4_custom_branding,
        example_5_tables_and_lists,
        example_6_info_boxes,
        example_7_invoice,
        example_8_multi_section
    ]

    for example_func in examples:
        try:
            example_func()
        except Exception as e:
            print(f"✗ Error in {example_func.__name__}: {e}")

    print("\n" + "="*70)
    print("All examples completed!")
    print("Check the data/exports/ directory for generated files.")
    print("="*70)


if __name__ == "__main__":
    from docx.shared import RGBColor  # Import here for example_4
    run_all_examples()
