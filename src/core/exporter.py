"""Export utilities for generating documents in various formats"""

import os
import re
from typing import Dict, Any
from pathlib import Path


class DocumentExporter:
    """Exporter for various document formats"""

    def __init__(self):
        """Initialize exporter"""
        pass

    def export(self, content: str, output_path: str, format: str = "txt", **kwargs) -> bool:
        """
        Generic export method that dispatches to appropriate format-specific method

        Args:
            content: Document content
            output_path: Output file path
            format: Output format (txt, md, html, pdf, docx)
            **kwargs: Additional format-specific arguments

        Returns:
            True if successful, False otherwise
        """
        format = format.lower()

        try:
            if format in ["txt", "text"]:
                self.export_to_text(content, output_path)
                return True
            elif format in ["md", "markdown"]:
                metadata = kwargs.get("metadata", None)
                self.export_to_markdown(content, output_path, metadata)
                return True
            elif format == "html":
                title = kwargs.get("title", "Document")
                self.export_to_html(content, output_path, title)
                return True
            elif format == "pdf":
                title = kwargs.get("title", "Document")
                return self.export_to_pdf(content, output_path, title)
            elif format == "docx":
                title = kwargs.get("title", "Document")
                return self.export_to_docx(content, output_path, title)
            else:
                raise ValueError(f"Unsupported format: {format}")
        except Exception as e:
            print(f"Export failed: {e}")
            return False

    def export_to_text(self, content: str, output_path: str) -> None:
        """
        Export to plain text file

        Args:
            content: Document content
            output_path: Output file path
        """
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def export_to_markdown(self, content: str, output_path: str, metadata: Dict[str, str] = None) -> None:
        """
        Export to Markdown file with optional frontmatter

        Args:
            content: Document content
            output_path: Output file path
            metadata: Optional metadata for frontmatter
        """
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            # Add frontmatter if metadata provided
            if metadata:
                f.write("---\n")
                for key, value in metadata.items():
                    f.write(f"{key}: {value}\n")
                f.write("---\n\n")

            f.write(content)

    def export_to_html(self, content: str, output_path: str, title: str = "Документ") -> None:
        """
        Export to HTML file

        Args:
            content: Document content
            output_path: Output file path
            title: Document title
        """
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

        # Convert to HTML (simple text to HTML with preserved formatting)
        html_content = self._text_to_html(content)

        html_template = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1, h2, h3 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        .variable {{
            background-color: #fff3cd;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: monospace;
        }}
        .block {{
            margin-bottom: 30px;
            padding: 20px;
            background-color: #f8f9fa;
            border-left: 4px solid #3498db;
        }}
        .section {{
            margin-bottom: 20px;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        @media print {{
            body {{
                max-width: 100%;
            }}
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)

    def export_to_pdf(self, content: str, output_path: str, title: str = "Документ") -> bool:
        """
        Export to PDF file (requires weasyprint or similar library)

        Args:
            content: Document content
            output_path: Output file path
            title: Document title

        Returns:
            True if successful, False otherwise
        """
        try:
            from weasyprint import HTML
            from io import BytesIO

            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

            # Create HTML first
            html_content = self._text_to_html(content)
            html_full = f"<html><head><meta charset='UTF-8'><title>{title}</title></head><body>{html_content}</body></html>"

            # Convert to PDF
            HTML(string=html_full).write_pdf(output_path)
            return True

        except ImportError:
            # WeasyPrint not installed, try alternative method
            try:
                # Create HTML file first
                html_path = output_path.replace('.pdf', '.html')
                self.export_to_html(content, html_path, title)
                return False  # PDF not created, but HTML created
            except Exception:
                return False

    def export_to_docx(self, content: str, output_path: str, title: str = "Документ") -> bool:
        """
        Export to DOCX file (requires python-docx library)

        Args:
            content: Document content
            output_path: Output file path
            title: Document title

        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

            doc = Document()

            # Add title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Process content
            lines = content.split('\n')
            for line in lines:
                line = line.strip()

                if not line:
                    continue

                # Detect headers
                if line.startswith('БЛОК') or line.startswith('0.'):
                    doc.add_heading(line, level=1)
                elif re.match(r'^\d+\.', line):
                    doc.add_heading(line, level=2)
                else:
                    # Regular paragraph
                    para = doc.add_paragraph()

                    # Highlight variables
                    parts = re.split(r'(\{[^}]+\})', line)
                    for part in parts:
                        if part.startswith('{') and part.endswith('}'):
                            run = para.add_run(part)
                            run.font.color.rgb = RGBColor(255, 140, 0)
                            run.bold = True
                        else:
                            para.add_run(part)

            doc.save(output_path)
            return True

        except ImportError:
            return False

    def _text_to_html(self, text: str) -> str:
        """
        Convert plain text to HTML with basic formatting

        Args:
            text: Plain text

        Returns:
            HTML string
        """
        lines = text.split('\n')
        html_parts = []
        in_list = False

        for line in lines:
            line_stripped = line.strip()

            if not line_stripped:
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append('<br>')
                continue

            # Headers
            if line_stripped.startswith('БЛОК') or line_stripped.startswith('МЕГА-ШАБЛОН'):
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append(f'<h1>{self._escape_html(line_stripped)}</h1>')
            elif re.match(r'^\d+\.', line_stripped):
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append(f'<h2>{self._escape_html(line_stripped)}</h2>')
            elif line_stripped.startswith('---'):
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append('<hr>')
            # Lists
            elif line_stripped.startswith('•') or line_stripped.startswith('-'):
                if not in_list:
                    html_parts.append('<ul>')
                    in_list = True
                html_parts.append(f'<li>{self._highlight_variables(line_stripped[1:].strip())}</li>')
            else:
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                html_parts.append(f'<p>{self._highlight_variables(line_stripped)}</p>')

        if in_list:
            html_parts.append('</ul>')

        return '\n'.join(html_parts)

    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters"""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))

    def _highlight_variables(self, text: str) -> str:
        """Highlight template variables in HTML"""
        text = self._escape_html(text)
        # Highlight {Variable} patterns
        text = re.sub(r'\{([^}]+)\}', r'<span class="variable">{\1}</span>', text)
        return text
